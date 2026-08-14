import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_shtpic_template(output_path: str):
    """
    Tự động tạo file Word template chuẩn SHTP-IC (.docx) với các biến Jinja2.
    """
    doc = Document()

    # Thiết lập căn lề chuẩn
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)

    # Tiêu đề Vườn ươm
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_header.add_run("VƯỜN ƯƠM DOANH NGHIỆP CÔNG NGHỆ CAO\n")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(13)
    r1.font.bold = True

    r2 = p_header.add_run("E2a-10 Đường D2b, Khu Công nghệ cao, P.Tăng Nhơn Phú, TP.HCM\nĐiện thoại: 028 7100 7986 - Website: www.shtpic.org\n")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(10)
    r2.font.italic = True

    doc.add_paragraph()

    # Tên Form
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("THUYẾT MINH DỰ ÁN\n")
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_sub = p_title.add_run("Tham gia chương trình ươm tạo tại Vườn ươm Doanh nghiệp Công nghệ cao\n")
    r_sub.font.name = "Times New Roman"
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True

    # Section I
    doc.add_heading("I. THÔNG TIN CHUNG VỀ DỰ ÁN", level=1)
    doc.add_paragraph("1. Tên dự án: {{ ten_du_an }}")
    doc.add_paragraph("2. Sản phẩm/dịch vụ chính: {{ san_pham_chinh }}")
    doc.add_paragraph("3. Công nghệ chính của dự án: {{ cong_nghe_chinh }}")
    doc.add_paragraph("4. Lĩnh vực: {{ linh_vuc }}")

    # Section II.1
    doc.add_heading("II. THÔNG TIN CHI TIẾT DỰ ÁN", level=1)
    doc.add_heading("1. Giới thiệu tóm tắt về dự án", level=2)
    doc.add_paragraph("- Vấn đề dự án đang muốn giải quyết và giải pháp dự án nêu ra:\n{{ van_de_giai_phap }}")
    doc.add_paragraph("- Tính cấp thiết thực hiện dự án:\n{{ tinh_cap_thiet }}")
    doc.add_paragraph("- Mục tiêu kinh tế - xã hội của dự án:\n{{ muc_tieu_ktxh }}")
    doc.add_paragraph("- Mục tiêu về khoa học công nghệ của dự án:\n{{ muc_tieu_khcn }}")

    # Section II.2
    doc.add_heading("2. Công nghệ, sản phẩm đăng ký ươm tạo của dự án", level=2)
    doc.add_heading("2.1. Sản phẩm đăng ký ươm tạo của dự án", level=3)
    doc.add_paragraph("Mô tả sản phẩm công nghệ cao đăng ký ươm tạo:\n- Giải trình minh chứng:\n{{ minh_chung_san_pham }}\n- Quy trình hoạt động & Cấu phần:\n{{ quy_trinh_san_pham }}\n- Chất lượng, tính năng vượt trội:\n{{ tinh_nang_vuot_troi }}")

    doc.add_heading("2.2. Công nghệ đăng ký ươm tạo của dự án", level=3)
    doc.add_paragraph("Hiện trạng ứng dụng, ươm tạo công nghệ cao:\n- Minh chứng công nghệ phù hợp:\n{{ minh_chung_cong_nghe }}\n- Quy trình công nghệ & Điểm nổi bật:\n{{ quy_trinh_cong_nghe }}\n- Tính khả thi kỹ thuật:\n{{ tinh_kha_thi_ky_thuat }}\n- Bảo mật & Khả năng mở rộng:\n{{ bao_mat_mo_rong }}")

    # Section II.5
    doc.add_heading("5. Tiềm năng thương mại & Phân tích đối thủ", level=2)
    doc.add_paragraph("- Tiềm năng thị trường:\n{{ tiem_nang_thi_truong }}")
    doc.add_paragraph("- Phân tích Đối thủ 1: {{ doi_thu_1 }}\n- Phân tích Đối thủ 2: {{ doi_thu_2 }}\n- Lợi thế cạnh tranh:\n{{ loi_the_canh_tranh }}")

    # Section III.2 Support Table
    doc.add_heading("III.2. CÁC NỘI DUNG DỰ ÁN ĐỀ XUẤT CẦN HỖ TRỢ", level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "TT"
    hdr_cells[1].text = "Thực trạng"
    hdr_cells[2].text = "Nhu cầu"
    hdr_cells[3].text = "Kết quả mong muốn"

    doc.add_paragraph("\n\n Tôi/Chúng tôi xin chịu trách nhiệm hoàn toàn về toàn bộ thông tin nêu trên.")

    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    create_shtpic_template("/Users/khoata/Documents/VƯ DN CNC/Code/shtp_proposal_ai/template_shtpic.docx")
    print("Created Word Template successfully!")
